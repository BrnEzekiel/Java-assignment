import asyncio
from markdown_it import MarkdownIt
from docx import Document
from pyppeteer import launch

assignments = [
    "assignment1_solutions/assignment1_report.md",
    "assignment2_solutions/assignment2_report.md"
]

async def md_to_pdf(md_path, pdf_path):
    with open(md_path, "r") as f:
        md_text = f.read()
    md = MarkdownIt()
    html = md.render(md_text)
    html = f"<html><body>{html}</body></html>"
    
    browser = await launch()
    page = await browser.newPage()
    await page.setContent(html)
    await page.pdf({'path': pdf_path, 'format': 'A4'})
    await browser.close()

def md_to_docx(md_path, docx_path):
    with open(md_path, "r") as f:
        md_text = f.read()
    doc = Document()
    for line in md_text.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    doc.save(docx_path)

async def main():
    for md_path in assignments:
        base = md_path.rsplit("/", 1)[0]
        name = md_path.rsplit("/", 1)[1].replace(".md", "")
        docx_path = f"{base}/{name}.docx"
        pdf_path = f"{base}/{name}.pdf"

        print(f"Generating {docx_path}...")
        md_to_docx(md_path, docx_path)
        print(f"Generating {pdf_path}...")
        await md_to_pdf(md_path, pdf_path)
        print(f"Done with {name}!\n")

asyncio.run(main())
